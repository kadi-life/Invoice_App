from django.http import HttpResponse
from django.template.loader import get_template
import os
from django.conf import settings
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO

def add_watermark(document, image_path):
    """Add a watermark to all pages of the document"""
    # This function is now disabled to remove the watermark/logo from the header
    # as requested by the user
    pass

def generate_invoice_docx(invoice):
    """Generate a Word document for an invoice"""
    document = Document()
    
    # Document properties
    document.core_properties.title = f"INVOICE {invoice.invoice_number.upper() if invoice.invoice_number else invoice.id}"
    document.core_properties.author = "Skids LOGISTICS LTD"
    
    # Company header in a box (like PDF)
    company_header = document.add_paragraph()
    company_header_run = company_header.add_run('NO. 17 Eastern Bypass, Buchi Atako Villa Port Harcourt, T: 07035495280, E: info@skidslogistics.com.')
    
    # Add border to the paragraph (box style)
    p = company_header._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        pBdr.append(border)
    pPr.append(pBdr)
    
    # Add logo after the company information
    try:
        logo_path = os.path.join(settings.STATIC_ROOT, 'images/skids_logo.png')
        if not os.path.exists(logo_path):
            # Fallback to old path if new path doesn't exist
            logo_path = os.path.join(settings.STATIC_ROOT, 'img/logo.png')
        if os.path.exists(logo_path):
            # Left align the logo (like PDF)
            logo_paragraph = document.add_paragraph()
            logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            logo_run = logo_paragraph.add_run()
            logo_run.add_picture(logo_path, width=Inches(3))
    except Exception as e:
        print(f"Error adding logo: {e}")
    
    # Add demarcation line
    demarcation = document.add_paragraph()
    demarcation_run = demarcation.add_run()
    demarcation_run.font.color.rgb = RGBColor(0, 0, 0)
    p = demarcation._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    
    # Add watermark if logo exists
    try:
        if os.path.exists(logo_path):
            add_watermark(document, logo_path)
    except Exception as e:
        print(f"Error adding watermark: {e}")
    
    # Invoice title (styled like PDF)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"INVOICE {invoice.invoice_number.upper() if invoice.invoice_number else invoice.id}")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.underline = True
    
    # Create a table for client and invoice information (like PDF layout)
    info_table = document.add_table(rows=1, cols=2)
    info_table.style = 'Table Grid'
    info_table.autofit = False
    
    # Set column widths
    for cell in info_table.columns[0].cells:
        cell.width = Inches(3.5)
    for cell in info_table.columns[1].cells:
        cell.width = Inches(3.5)
    
    # Client information (left column)
    client_cell = info_table.cell(0, 0)
    client_info = client_cell.paragraphs[0]
    client_info.add_run('CLIENT: ').bold = True
    client_info.add_run(f"{invoice.client_name}")
    
    # Invoice information (right column)
    invoice_cell = info_table.cell(0, 1)
    invoice_info = invoice_cell.paragraphs[0]
    invoice_info.add_run('DATE: ').bold = True
    invoice_info.add_run(f"{invoice.date_created.strftime('%d-%m-%Y')}")
    invoice_info.add_run('\nDUE DATE: ').bold = True
    invoice_info.add_run(f"{invoice.due_date.strftime('%d-%m-%Y')}")
    invoice_info.add_run('\nSTATUS: ').bold = True
    invoice_info.add_run(f"{invoice.status}")

    
    # Items table
    document.add_paragraph()
    items_table = document.add_table(rows=1, cols=5)
    items_table.style = 'Table Grid'
    
    # Header row
    header_cells = items_table.rows[0].cells
    header_cells[0].text = "DESCRIPTION"
    header_cells[1].text = "QUANTITY"
    header_cells[2].text = f"UNIT PRICE ({invoice.currency})"
    header_cells[3].text = f"TOTAL ({invoice.currency})"
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add items
    for item in invoice.items.all():
        row_cells = items_table.add_row().cells
        row_cells[0].text = item.name
        # Extract numeric part from quantity and add unit if available
        quantity_text = str(item.quantity)
        if hasattr(item, 'unit') and item.unit:
            quantity_text += f" {item.unit}"
        row_cells[1].text = quantity_text
        # Format numbers with comma delimiter
        price = '{:,.0f}'.format(item.price)
        total = '{:,.0f}'.format(item.total)
        row_cells[2].text = price
        row_cells[3].text = total
        row_cells[4].text = item.lead_time if item.lead_time else "1-7 DAYS"
    
    # Totals
    document.add_paragraph()
    
    # Create a left-aligned paragraph for each total line to match PDF style
    subtotal_para = document.add_paragraph()
    subtotal_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtotal_para.add_run(f"Subtotal: {invoice.currency} {'{:,.0f}'.format(invoice.subtotal)}")
    
    vat_para = document.add_paragraph()
    vat_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    vat_para.add_run(f"VAT ({invoice.vat_percentage}%): {invoice.currency} {'{:,.0f}'.format(invoice.vat_amount)}")
    
    # Add a line before the total
    document.add_paragraph().add_run('_' * 30)
    
    # Make the total bold
    total_para = document.add_paragraph()
    total_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    total_run = total_para.add_run(f"Total: {invoice.currency} {'{:,.0f}'.format(invoice.total)}")
    total_run.bold = True
    
    # Notes if present
    if invoice.notes:
        document.add_paragraph()
        notes_para = document.add_paragraph()
        notes_run = notes_para.add_run("Notes:")
        notes_run.bold = True
        document.add_paragraph(invoice.notes)
    
    # Add images if any
    if any(item.image for item in invoice.items.all()):
        document.add_page_break()
        
        # Add images title
        images_title = document.add_paragraph()
        images_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        images_title_run = images_title.add_run("ITEM IMAGES")
        images_title_run.bold = True
        images_title_run.font.size = Pt(14)
        
        # Create a 2x2 table for images
        items_with_images = [item for item in invoice.items.all() if item.image]
        
        # Process 4 images per page
        for i in range(0, len(items_with_images), 4):
            # Create a 2x2 table
            if i > 0:
                document.add_page_break()
                # Add images title on new page
                new_page_title = document.add_paragraph()
                new_page_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                new_page_title_run = new_page_title.add_run("ITEM IMAGES (CONTINUED)")
                new_page_title_run.bold = True
                new_page_title_run.font.size = Pt(14)
            
            table = document.add_table(rows=2, cols=2)
            table.style = 'Table Grid'
            
            # Add images to the table
            for j in range(4):
                if i + j < len(items_with_images):
                    item = items_with_images[i + j]
                    row = j // 2
                    col = j % 2
                    
                    cell = table.cell(row, col)
                    paragraph = cell.paragraphs[0]
                    
                    # Add image
                    try:
                        image_path = item.image.path
                        print(f"\n\n[DEBUG DOCX] Image path: {image_path} (Exists: {os.path.exists(image_path)})\n\n")
                        
                        # If the image path doesn't exist, try to find it in the media root
                        if not os.path.exists(image_path):
                            # Try to get the image URL and convert it to a path
                            try:
                                image_url = item.image.url
                                print(f"[DEBUG DOCX] Image URL: {image_url}")
                                
                                # Convert URL to path
                                if image_url.startswith(settings.MEDIA_URL):
                                    rel_path = image_url.replace(settings.MEDIA_URL, '')
                                    alt_path = os.path.join(settings.MEDIA_ROOT, rel_path)
                                    print(f"[DEBUG DOCX] Alternative path from URL: {alt_path} (Exists: {os.path.exists(alt_path)})")
                                    if os.path.exists(alt_path):
                                        image_path = alt_path
                                        
                                # Try alternative path in item_images directory
                                if not os.path.exists(image_path):
                                    image_filename = os.path.basename(str(item.image))
                                    alternative_path = os.path.join(settings.MEDIA_ROOT, 'item_images', image_filename)
                                    print(f"[DEBUG DOCX] Trying alternative path in item_images: {alternative_path}")
                                    if os.path.exists(alternative_path):
                                        image_path = alternative_path
                            except Exception as url_error:
                                print(f"[DEBUG DOCX] Error getting image URL: {url_error}")
                        
                        # If we have a valid image path, add it to the document
                        if os.path.exists(image_path):
                            run = paragraph.add_run()
                            run.add_picture(image_path, width=Inches(2.5))
                            
                            # Add item name below image
                            name_paragraph = cell.add_paragraph()
                            name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            name_run = name_paragraph.add_run(item.name)
                            name_run.bold = True
                        else:
                            print(f"[DEBUG DOCX] Image file not found: {image_path}")
                            paragraph.add_run(f"[Image not available for {item.name}]")
                            
                            # Add item name even if image is missing
                            name_paragraph = cell.add_paragraph()
                            name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            name_run = name_paragraph.add_run(f"[No Image] {item.name}")
                            name_run.bold = True
                    except Exception as e:
                        print(f"[DEBUG DOCX] Error adding image: {e}")
                        paragraph.add_run(f"[Image not available for {item.name}]")
                        
                        # Add item name even if there's an error
                        name_paragraph = cell.add_paragraph()
                        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        name_run = name_paragraph.add_run(f"[Image Error] {item.name}")
                        name_run.bold = True
    
    # Save to memory
    docx_buffer = BytesIO()
    document.save(docx_buffer)
    docx_buffer.seek(0)
    
    # Create response
    response = HttpResponse(docx_buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="Invoice_{invoice.invoice_number}.docx"'
    
    return response

def generate_quotation_docx(quotation, request=None):
    """Generate a Word document for a quotation"""
    document = Document()
    
    # Print debug information about the request and media paths
    if request:
        print(f"[DEBUG DOCX] Request scheme: {request.scheme}")
        print(f"[DEBUG DOCX] Request host: {request.get_host()}")
        print(f"[DEBUG DOCX] MEDIA_ROOT: {settings.MEDIA_ROOT}")
        print(f"[DEBUG DOCX] MEDIA_URL: {settings.MEDIA_URL}")
    
    # Document properties
    document.core_properties.title = f"Quotation Number {quotation.quotation_number.upper() if quotation.quotation_number else quotation.id}"
    document.core_properties.author = "Skids LOGISTICS LTD"
    
    # Company header in a box (like PDF)
    company_header = document.add_paragraph()
    company_header_run = company_header.add_run('NO. 17 Eastern Bypass, Buchi Atako Villa Port Harcourt, T: 07035495280, E: info@skidslogistics.com.')
    
    # Add border to the paragraph (box style)
    p = company_header._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        pBdr.append(border)
    pPr.append(pBdr)
    
    # Add logo after the company information
    try:
        logo_path = os.path.join(settings.STATIC_ROOT, 'images/skids_logo.png')
        if not os.path.exists(logo_path):
            # Fallback to old path if new path doesn't exist
            logo_path = os.path.join(settings.STATIC_ROOT, 'img/logo.png')
        if os.path.exists(logo_path):
            # Left align the logo (like PDF)
            logo_paragraph = document.add_paragraph()
            logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            logo_run = logo_paragraph.add_run()
            logo_run.add_picture(logo_path, width=Inches(3))
    except Exception as e:
        print(f"Error adding logo: {e}")
    
    # Add demarcation line
    demarcation = document.add_paragraph()
    demarcation_run = demarcation.add_run()
    demarcation_run.font.color.rgb = RGBColor(0, 0, 0)
    p = demarcation._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    
    # Add watermark if logo exists
    try:
        if os.path.exists(logo_path):
            add_watermark(document, logo_path)
    except Exception as e:
        print(f"Error adding watermark: {e}")
    
    # Quotation title (styled like PDF)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"QUOTATION {quotation.quotation_number.upper() if quotation.quotation_number else quotation.id}")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.underline = True
    
    # Create a table for client and quotation information (like PDF layout)
    info_table = document.add_table(rows=1, cols=2)
    info_table.style = 'Table Grid'
    info_table.autofit = False
    
    # Set column widths
    for cell in info_table.columns[0].cells:
        cell.width = Inches(3.5)
    for cell in info_table.columns[1].cells:
        cell.width = Inches(3.5)
    
    # Client information (left column)
    client_cell = info_table.cell(0, 0)
    client_info = client_cell.paragraphs[0]
    client_info.add_run('CLIENT: ').bold = True
    client_info.add_run(f"{quotation.client_name}")
    
    # Quotation information (right column)
    quotation_cell = info_table.cell(0, 1)
    quotation_info = quotation_cell.paragraphs[0]
    quotation_info.add_run('DATE: ').bold = True
    quotation_info.add_run(f"{quotation.date_created.strftime('%d-%m-%Y')}")
    quotation_info.add_run('\nVALIDITY: ').bold = True
    quotation_info.add_run("30 days from date of issue")
    
    # Items table
    document.add_paragraph()
    items_table = document.add_table(rows=1, cols=5)
    items_table.style = 'Table Grid'
    
    # Header row
    header_cells = items_table.rows[0].cells
    header_cells[0].text = "DESCRIPTION"
    header_cells[1].text = "QUANTITY"
    header_cells[2].text = f"UNIT PRICE ({quotation.currency})"
    header_cells[3].text = f"TOTAL ({quotation.currency})"
    header_cells[4].text = "LEAD TIME"
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add items
    for item in quotation.items.all():
        row_cells = items_table.add_row().cells
        row_cells[0].text = item.name
        # Extract numeric part from quantity and add unit if available
        quantity_text = str(item.quantity)
        if hasattr(item, 'unit') and item.unit:
            quantity_text += f" {item.unit}"
        row_cells[1].text = quantity_text
        # Format numbers with comma delimiter
        price = '{:,.0f}'.format(item.price)
        total = '{:,.0f}'.format(item.total)
        row_cells[2].text = price
        row_cells[3].text = total
        row_cells[4].text = item.lead_time if item.lead_time else "1-7 DAYS"
    
    # Totals
    document.add_paragraph()
    
    # Create a left-aligned paragraph for each total line to match PDF style
    subtotal_para = document.add_paragraph()
    subtotal_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtotal_para.add_run(f"Subtotal: {quotation.currency} {'{:,.0f}'.format(quotation.subtotal)}")
    
    vat_para = document.add_paragraph()
    vat_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    vat_para.add_run(f"VAT ({quotation.vat_percentage}%): {quotation.currency} {'{:,.0f}'.format(quotation.vat_amount)}")
    
    # Add a line before the total
    document.add_paragraph().add_run('_' * 30)
    
    # Make the total bold
    total_para = document.add_paragraph()
    total_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    total_run = total_para.add_run(f"Total: {quotation.currency} {'{:,.0f}'.format(quotation.total)}")
    total_run.bold = True
    
    # Notes if present
    if quotation.notes:
        document.add_paragraph()
        notes_para = document.add_paragraph()
        notes_run = notes_para.add_run("Notes:")
        notes_run.bold = True
        document.add_paragraph(quotation.notes)
    
    # Add images if any
    if any(item.image for item in quotation.items.all()):
        document.add_page_break()
        
        # Add images title
        images_title = document.add_paragraph()
        images_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        images_title_run = images_title.add_run("ITEM IMAGES")
        images_title_run.bold = True
        images_title_run.font.size = Pt(14)
        
        # Create a 2x2 table for images
        items_with_images = [item for item in quotation.items.all() if item.image]
        
        # Process 4 images per page
        for i in range(0, len(items_with_images), 4):
            # Create a 2x2 table
            if i > 0:
                document.add_page_break()
                # Add images title on new page
                new_page_title = document.add_paragraph()
                new_page_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                new_page_title_run = new_page_title.add_run("ITEM IMAGES (CONTINUED)")
                new_page_title_run.bold = True
                new_page_title_run.font.size = Pt(14)
            
            table = document.add_table(rows=2, cols=2)
            table.style = 'Table Grid'
            
            # Add images to the table
            for j in range(4):
                if i + j < len(items_with_images):
                    item = items_with_images[i + j]
                    row = j // 2
                    col = j % 2
                    
                    cell = table.cell(row, col)
                    paragraph = cell.paragraphs[0]
                    
                    # Add image
                    try:
                        image_path = item.image.path
                        print(f"\n\n[DEBUG DOCX] Quotation Image path: {image_path} (Exists: {os.path.exists(image_path)})\n\n")
                        
                        # If the image path doesn't exist, try to find it in the media root
                        if not os.path.exists(image_path):
                            # Try to get the image URL and convert it to a path
                            try:
                                image_url = item.image.url
                                print(f"[DEBUG DOCX] Quotation Image URL: {image_url}")
                                
                                # Convert URL to path
                                if image_url.startswith(settings.MEDIA_URL):
                                    rel_path = image_url.replace(settings.MEDIA_URL, '')
                                    alt_path = os.path.join(settings.MEDIA_ROOT, rel_path)
                                    print(f"[DEBUG DOCX] Quotation Alternative path from URL: {alt_path} (Exists: {os.path.exists(alt_path)})")
                                    if os.path.exists(alt_path):
                                        image_path = alt_path
                                        
                                # Try alternative path in item_images directory
                                if not os.path.exists(image_path):
                                    image_filename = os.path.basename(str(item.image))
                                    alternative_path = os.path.join(settings.MEDIA_ROOT, 'item_images', image_filename)
                                    print(f"[DEBUG DOCX] Trying alternative path in item_images: {alternative_path}")
                                    if os.path.exists(alternative_path):
                                        image_path = alternative_path
                            except Exception as url_error:
                                print(f"[DEBUG DOCX] Error getting quotation image URL: {url_error}")
                        
                        # If we have a valid image path, add it to the document
                        if os.path.exists(image_path):
                            run = paragraph.add_run()
                            run.add_picture(image_path, width=Inches(2.5))
                            
                            # Add item name below image
                            name_paragraph = cell.add_paragraph()
                            name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            name_run = name_paragraph.add_run(item.name)
                            name_run.bold = True
                        else:
                            print(f"[DEBUG DOCX] Quotation Image file not found: {image_path}")
                            paragraph.add_run(f"[Image not available for {item.name}]")
                            
                            # Add item name even if image is missing
                            name_paragraph = cell.add_paragraph()
                            name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            name_run = name_paragraph.add_run(f"[No Image] {item.name}")
                            name_run.bold = True
                    except Exception as e:
                        print(f"[DEBUG DOCX] Error adding quotation image: {e}")
                        paragraph.add_run(f"[Image not available for {item.name}]")
                        
                        # Add item name even if there's an error
                        name_paragraph = cell.add_paragraph()
                        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        name_run = name_paragraph.add_run(f"[Image Error] {item.name}")
                        name_run.bold = True
    
    # Save to memory
    docx_buffer = BytesIO()
    document.save(docx_buffer)
    docx_buffer.seek(0)
    
    # Create response
    response = HttpResponse(docx_buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="Quotation_{quotation.quotation_number}.docx"'
    
    return response